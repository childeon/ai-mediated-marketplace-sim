"""
make_report.py — Generate report_final.docx from the simulation results.
Run: python3 make_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────
BLUE = RGBColor(0x1F, 0x45, 0x78)   # dark academic blue

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13 if level == 1 else 11, bold=True, color=BLUE)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True)
    return p

def body(text, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def body_parts(parts, space_after=4):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, size=11, bold=bold, italic=italic)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=11, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size=11)
    else:
        run = p.add_run(text)
        set_font(run, size=11)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=10, bold=True, italic=True)
    return p

def add_table(header_row, data_rows, bold_header=True):
    n_cols = len(header_row)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = tbl.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(cell_text)
        set_font(run, size=10, bold=bold_header)
        # Blue header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4578")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=10)
            # Zebra striping
            if r_idx % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EBF0F7")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)

    doc.add_paragraph()  # spacer
    return tbl

def monospace_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after  = Pt(6)
r = p_title.add_run("AI-Mediated Food Delivery Marketplace Simulation:")
set_font(r, size=16, bold=True, color=BLUE)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(10)
r = p_sub.add_run("Monetisation Design, Sponsorship Bias, and Consumer Welfare")
set_font(r, size=14, bold=True, color=BLUE)

for line in [
    "Columbia University · Department of Industrial Engineering & Operations Research",
    "Digital Marketplaces · Spring 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    set_font(r, size=11, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

heading("1. Introduction")
body(
    "Food delivery platforms such as DoorDash, Uber Eats, and Grubhub operate as two-sided platforms, "
    "connecting merchants (restaurants), consumers, and logistics networks (couriers) simultaneously. "
    "They benefit from network effects and economies of scale, monetising through commissions, "
    "advertising, and subscription programs."
)
body(
    "However, OpenAI's entry into food delivery has begun to reshape this landscape. Its partnership "
    "with Uber allows users to receive food recommendations and complete transactions directly through "
    "ChatGPT. As a result, the industry is now debating whether LLMs could evolve into super apps that "
    "capture the majority of user traffic and advertising value. We believe food delivery platforms still "
    "retain unique value through fulfilment and service delivery, but it is equally important to "
    "understand the new advertising dynamics created by this shift."
)
body(
    "Our research focuses on how much platforms should be willing to pay for advertising in order for "
    "their restaurants to be recommended by an LLM — and, critically, whether the design of that "
    "payment mechanism determines whether LLM intermediation is welfare-improving or welfare-degrading."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════

heading("2. Workflow of Food Delivery")
body(
    "The traditional food delivery decision-making process follows a search–pick–order flow. "
    "Restaurants are typically recommended based on keyword matching, along with advertising budgets "
    "and sponsored placements. With AI-assisted food delivery, however, customers can enter "
    "natural-language prompts, and the LLM can recommend the top restaurants that best match the "
    "user's intent while also taking into account factors such as price, delivery time, quality, and "
    "advertising exposure."
)
body(
    "As illustrated by real-world cases from ChatGPT and Qwen, the AI-guided decision process follows "
    "a structured path: the user expresses a need (cuisine, budget, time constraint) → the LLM "
    "aggregates options across platforms → it returns a ranked shortlist → the user clicks through or "
    "ignores the recommendation → the platform completes the transaction."
)
body(
    "This flow compresses the traditional search–compare–decide process into a guided interaction. For "
    "users, it reduces search costs and decision time. For platforms and merchants, it improves "
    "conversion by narrowing the consideration set. At the same time, AI intermediates traffic "
    "allocation, shaping merchant visibility and demand distribution within the marketplace. Crucially, "
    "who pays the AI — and how — determines whether this traffic allocation serves consumers or distorts "
    "the market in favour of the highest bidder."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. QUALITATIVE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Qualitative Analysis of Equilibrium")

subheading("3.1  Impact of AI on Food Delivery Platforms: Three Pillars")
body(
    "We believe traditional food delivery platforms will remain necessary because LLMs cannot execute "
    "physical delivery (riders, logistics networks). However, the AI layer has significant implications "
    "for all three platform pillars."
)

body_parts([
    ("On thickness, ", True, False),
    ("LLMs can materially increase demand aggregation because they sit at the top of the funnel where "
     "consumers express intent in natural language, often with far larger weekly active user bases than "
     "any single delivery app. By reducing search costs and helping users articulate preferences — "
     "budget, cuisine, dietary restrictions, delivery speed — AI can convert more latent demand into "
     "actual orders and expand the overall volume flowing into food delivery ecosystems. However, this "
     "thicker market does not necessarily benefit Uber Eats or DoorDash directly. AI may shift demand "
     "ownership upward to the LLM layer, turning delivery platforms into downstream fulfilment providers "
     "rather than the primary destination where users begin discovery. In our simulation, the neutral "
     "LLM raises overall conversion (from 80.4% to 81.6%) and improves intent match quality — but does "
     "so while simultaneously reducing aggregate platform net revenue, because better matching routes "
     "consumers away from high-commission platform-promoted offers toward better-fit, lower-sponsored "
     "restaurants.", False, False),
])

body_parts([
    ("On trust, ", True, False),
    ("AI is much more limited. A conversational agent can recommend restaurants, summarise reviews, or "
     "optimise choices, but it cannot guarantee food quality, delivery accuracy, refund resolution, or "
     "service reliability. Moreover, because AI-mediated ranking is less transparent than traditional "
     "app search results, it intensifies the tension between user relevance and monetisation. The "
     "delivery platforms retain a core role as the execution and accountability layer, owning logistics, "
     "dispatch, payment rails, and customer support. Our simulation operationalises this by preserving "
     "a hard cuisine-match criterion in intent fulfilment: the LLM earns no credit for recommending a "
     "wrong-cuisine option regardless of how well it fits on price or time.", False, False),
])

body_parts([
    ("On value add, ", True, False),
    ("LLMs create real consumer value not by replacing delivery infrastructure but by improving "
     "decision-making. Food delivery is a high-friction use case characterised by excessive choice, "
     "repeated browsing, and weak personalisation. The LLM acts as a more intelligent interface that "
     "reduces decision fatigue, provides contextual recommendations, and compares options across "
     "multiple apps. In our simulation, the LLM's key value is cross-platform aggregation — it is the "
     "only mechanism that eliminates the 'platform sampling problem' where consumers, browsing only "
     "their 2 most-loyal apps, may entirely miss the cuisine they want. The strategic question is thus "
     "not whether AI replaces food delivery platforms, but whether those platforms can preserve "
     "ownership of the customer relationship in a world where discovery is intermediated by a "
     "general-purpose AI assistant.", False, False),
])

subheading("3.2  Qualitative Analysis of the New Equilibrium")
body(
    "There are three key parties in AI-supported food delivery: food delivery platforms (in our model, "
    "QuickEats, FoodRush, and NomNom), consumers, and the LLM intermediary. The introduction of the "
    "LLM layer changes the traditional advertising equilibrium by inserting a new gatekeeper between "
    "consumer intent and platform fulfilment."
)

body_parts([
    ("For platforms, ", True, False),
    ("paying for recommendation exposure at the LLM can be viewed as a new form of customer "
     "acquisition cost. However, neutral LLM access actually generates a negative willingness-to-pay "
     "in our calibrated simulation (−$1,483 aggregate net vs. direct search). Paid regimes restore "
     "positive WTP only because routing bias concentrates orders in higher-commission platforms, "
     "creating a gross surplus the LLM then captures.", False, False),
])

body_parts([
    ("For customers, ", True, False),
    ("the LLM creates value by reducing search and comparison costs. The LLM's cross-platform "
     "aggregation raises intent fulfilment from 0.683 to 0.747 and lowers average paid price by "
     "$1.71. However, under strong sponsorship bias, these consumer gains erode — and at bias "
     "strength ≥ 1.0, the LLM-mediated result is actually worse than loyal direct search on "
     "intent fulfilment (0.649 vs. 0.683).", False, False),
])

body_parts([
    ("For the LLM operator, ", True, False),
    ("advertising is a potential monetisation lever, but only if user trust remains intact. The "
     "CPFI regime resolves this trade-off most cleanly — paying the LLM only when intent is "
     "fulfilled aligns its financial incentive with the consumer's welfare, achieving the best "
     "shortlist relevance (0.981) among paid regimes.", False, False),
])

body(
    "The new equilibrium therefore depends on balancing these incentives. Food delivery platforms "
    "will participate in LLM advertising not because they fully control this new channel, but because "
    "they cannot afford to be absent from a discovery layer that increasingly shapes consumer choice."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════

heading("4. Methodology")

subheading("4.1  System Architecture")
body(
    "The simulation implements a five-module architecture. World construction (entities.py) instantiates "
    "three platforms — QuickEats, FoodRush, and NomNom — 50 restaurants (each listed on 1–3 platforms), "
    "restaurant–platform offer tuples with quality–price correlation (ρ ≈ 0.5), and 1,000 consumers "
    "drawn from configured distributions using a seeded RNG. Restaurant quality is snapped to half-star "
    "precision (2.0–5.0 ★), with menu prices ranging from $19.00 to $50.00 positively correlated with "
    "quality."
)
body(
    "Ranking logic (ranking.py) maintains two scoring systems — platform-internal ranking and LLM "
    "relevance scoring — together with a 9×9 symmetric cuisine similarity table covering Italian, "
    "Chinese, Indian, Mexican, Japanese, American, Thai, Korean, and Mediterranean cuisines. Consumer "
    "choice (choice.py) implements a two-stage LLM interaction model: a click stage and a post-click "
    "purchase stage, each with its own outside option, plus a single-stage direct choice for the "
    "no-LLM world. The simulation engine (simulation.py) orchestrates consumer sessions and records "
    "per-transaction outcomes including LLM payments, click behaviour, and platform revenue. Metrics "
    "aggregation (metrics.py) computes consumer, restaurant, platform, and LLM-side statistics "
    "including Gini and HHI concentration measures. An attribution module (attribution.py) performs "
    "counterfactual causal analysis via Monte Carlo re-simulation."
)

subheading("4.2  Experimental Worlds")
body_parts([
    ("World A (No-LLM) ", True, False),
    ("models a loyalty-aware fragmented search. Each consumer identifies the 2 platforms to which they "
     "are most loyal, applies a cuisine filter (searching for their desired cuisine on each app), and "
     "observes the platform's top-5 ranked offers within that cuisine on each platform — a pool of up "
     "to 10 offers. Selection then proceeds via multinomial logit over the observed pool plus an outside "
     "option (utility = 2.5, representing not ordering).", False, False),
])
body_parts([
    ("World B (LLM-Mediated) ", True, False),
    ("models AI-brokered discovery with a two-stage interaction. The consumer submits a natural-language "
     "intent — cuisine preference, budget, maximum delivery time — to the LLM, which scores all offers "
     "across all platforms simultaneously, optionally applies a sponsorship bias term, and returns a "
     "top-K = 5 shortlist. The consumer then decides whether to click one recommendation (outside "
     "option utility = 1.2 for ignoring the shortlist) and, if clicked, whether to complete the "
     "purchase (outside option utility = 0.6 for abandonment). CPC payments are triggered at the click "
     "stage; CPA and CPFI payments require a completed purchase.", False, False),
])

subheading("4.3  Scoring Formulas")
body(
    "Platform-internal scoring weights quality (25%, normalised to 5 stars), commission revenue "
    "(25%), promo discount boost (15%), negative total price (15%, normalised at $40), negative "
    "delivery time (10%, normalised at 90 min), and a sponsored ad boost (10%)."
)
body(
    "LLM relevance scoring prioritises semantic cuisine match (35%), quality (28%), affordability "
    "fit (20% — a graded function from 1.0 if within budget, declining linearly beyond), negative "
    "delivery time (12%), and negative total price (5%). Under paid regimes, a sponsorship bias "
    "term is added additively:"
)
bullet("CPC / CPA:  bias_term = bias_strength × (platform_spend / max_spend)", bold_prefix=None)
bullet(
    "CPFI:  bias_term = bias_strength × (platform_spend / max_spend) × intent_fulfilment(offer, intent)  "
    "— bias is modulated by how well the offer satisfies the consumer's intent, so off-intent "
    "sponsored offers receive reduced amplification."
)

body("Semantic match is continuous via the 9×9 cuisine similarity table (1.0 = exact; 0.05 = default for unlisted pairs). Consumer utility is:")

monospace_block([
    "U = 2.0 × cuisine_match",
    "  + 1.5 × quality_sens × (quality / 5)",
    "  − 1.2 × price_sens  × (total_price / 40)",
    "  − 0.8 × time_sens   × (delivery_time / 90)",
    "  + 0.5 × (promo_discount / 5)",
    "  + 0.6 × llm_trust × cuisine_match",
    "  + 0.4 × platform_loyalty",
    "  + ε,    ε ~ Normal(0, 0.3)",
])

body(
    "Sensitivity parameters (price, time, quality) are drawn per consumer from truncated normal "
    "distributions. The LLM trust bonus activates only when the offered cuisine matches the "
    "consumer's intent. Platform loyalty governs which platforms the consumer checks in the "
    "no-LLM world."
)

subheading("4.4  Intent Fulfilment Metric")
body(
    "The primary consumer welfare metric is intent fulfilment (0–1). A wrong cuisine yields a "
    "hard zero. Budget overrun is penalised by 0.5 per fractional dollar over budget (as a "
    "fraction of budget). Time overrun is penalised by 0.02 per minute. Quality scales the "
    "score from 0.60 (at 2★) to 1.00 (at 5★), reflecting the real cost of being matched to a "
    "low-quality restaurant even when cuisine, price, and time constraints are met. This metric "
    "measures whether a transaction satisfied the consumer's stated request independently of "
    "subjective preference."
)

subheading("4.5  Experimental Design")
caption("Table 1. Experimental design summary.")
add_table(
    ["Experiment", "Variable", "Fixed Condition", "Core Question"],
    [
        ["1", "World (No-LLM vs. LLM Neutral)", "Neutral regime",
         "Does LLM intermediation improve on loyalty-based search?"],
        ["2", "Regime (Neutral / CPC / CPA / CPFI)", "LLM world",
         "How does payment structure affect outcomes and platform WTP?"],
        ["3", "Bias strength (0.0 → 1.0)", "CPC regime",
         "How much bias is tolerable before LLM is worse than no-LLM?"],
        ["4", "Consumer trust range", "LLM Neutral",
         "Does trust adoption shift market power and match quality?"],
    ]
)
body(
    "Each experiment runs 10 Monte Carlo iterations of 1,000 consumers (10,000 sessions per "
    "scenario). For Experiment 2, paid regime rates are calibrated endogenously: the model "
    "first estimates each regime's gross platform surplus with routing bias active but LLM "
    "payments set to zero, then scales actual payments to 50% of aggregate gross positive "
    "surplus. A separate attribution module performs counterfactual causal analysis on a "
    "30-consumer sample with 50 re-simulations per consumer."
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. RESULTS
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Results")

subheading("5.1  Experiment 1: LLM vs. No-LLM Baseline")
body(
    "The neutral LLM improves consumer welfare compared to the loyalty-aware baseline, though the "
    "improvement is more modest than a naive cross-platform aggregation would suggest. Intent "
    "fulfilment rises from 0.683 to 0.747 (+9.3%), average price paid falls by $1.71, and average "
    "delivery time falls by 1.7 minutes. Conversion increases 1.2 percentage points."
)

caption("Table 2. Experiment 1 results: No-LLM vs. LLM Neutral (10 runs × 1,000 consumers).")
add_table(
    ["Metric", "No-LLM", "LLM Neutral", "Change"],
    [
        ["Intent fulfilment",       "0.683",   "0.747",   "+9.3%"],
        ["Consumer utility",        "2.368",   "2.375",   "+0.3%"],
        ["Conversion rate",         "80.4%",   "81.6%",   "+1.2 pp"],
        ["Avg price paid",          "$34.19",  "$32.48",  "−$1.71"],
        ["Avg delivery time",       "36.4 min","34.7 min","−1.7 min"],
        ["Shortlist relevance",     "—",       "0.972",   "—"],
        ["LLM click-through rate",  "—",       "93.9%",   "—"],
        ["Platforms inspected",     "2.0",     "1.0",     "−50% friction"],
        ["Restaurant profit/order", "$11.45",  "$10.47",  "−8.6%"],
        ["Platform net revenue",    "$69,366", "$67,884", "−2.1%"],
        ["Restaurant exposure HHI", "0.0202",  "0.0215",  "+6.4%"],
    ]
)

body(
    "Three mechanisms drive the intent fulfilment gain. First, the LLM aggregates all three "
    "platforms, eliminating the risk that a consumer's two most-loyal platforms both lack the "
    "desired cuisine. Second, the LLM's scoring simultaneously balances cuisine match, "
    "affordability, quality, and delivery time. Third, the LLM applies a quality adjustment "
    "in its intent score, surfacing higher-rated options within the consumer's budget."
)
body(
    "Critically, the neutral LLM generates a negative platform willingness-to-pay (−$1,483 "
    "aggregate net vs. no-LLM). By routing consumers to the best-fitting restaurants across "
    "all platforms — which are often lower-commission or lower-sponsored — the LLM reduces "
    "platform net revenue relative to a world where loyal consumers browse the "
    "highest-commission platforms first. Restaurant HHI increases slightly (+6.4%), since the "
    "LLM concentrates attention on the same high-fit restaurants across all sessions."
)

subheading("5.2  Experiment 2: Monetisation Regimes")
body(
    "Introducing payment structures re-establishes positive platform WTP in all paid regimes, "
    "but at different costs to consumer welfare and market structure. All paid regime payments "
    "are calibrated to 50% of the regime-specific aggregate gross platform surplus."
)

caption("Table 3. Experiment 2 results across monetisation regimes.")
add_table(
    ["Metric", "Neutral", "CPC", "CPA", "CPFI"],
    [
        ["Intent fulfilment",       "0.747",   "0.716",   "0.726",   "0.747"],
        ["Shortlist relevance",     "0.973",   "0.929",   "0.946",   "0.981"],
        ["Consumer utility",        "2.375",   "2.310",   "2.331",   "2.363"],
        ["Avg total price",         "$32.48",  "$32.23",  "$32.21",  "$32.83"],
        ["LLM revenue",             "$0",      "$3,101",  "$2,873",  "$1,942"],
        ["Effective LLM rate/order","$0.00",   "$0.39",   "$0.36",   "$0.24"],
        ["Platform net revenue",    "$67,884", "$72,468", "$72,239", "$71,309"],
        ["Platform net vs. no-LLM", "−$1,483", "+$3,101", "+$2,873", "+$1,942"],
        ["FoodRush order share",    "32.6%",   "66.7%",   "60.4%",   "46.5%"],
        ["NomNom order share",      "26.9%",   "3.8%",    "3.6%",    "12.8%"],
    ]
)

body_parts([
    ("CPC causes the most severe consumer harm and market distortion. ", True, False),
    ("FoodRush, which pays the highest CPC rate ($0.80/click), jumps from 33% to 67% of orders. "
     "NomNom — which pays the lowest rate ($0.30) — is nearly eliminated (3.8%). The mechanism "
     "is structural: per-click payments incentivise the LLM to maximise impressions for the "
     "highest-paying platform regardless of whether those impressions result in completed orders "
     "or satisfied intents.", False, False),
])

body_parts([
    ("CPFI is the most consumer-aligned paid regime. ", True, False),
    ("Its intent fulfilment (0.747) and shortlist relevance (0.981) are indistinguishable from "
     "or better than neutral. The CPFI bias mechanism multiplies each offer's sponsorship boost "
     "by its predicted intent fulfilment score, so off-intent sponsored offers receive little "
     "amplification while on-intent offers receive full amplification. This aligns the LLM's "
     "sponsorship incentive with consumer welfare. CPFI generates the lowest LLM revenue "
     "($1,942) among paid regimes, reflecting a genuine ceiling on revenue when payments require "
     "outcome satisfaction.", False, False),
])

body(
    "All paid regimes generate positive aggregate platform net revenue vs. no-LLM (unlike neutral), "
    "because routing bias concentrates orders in higher-commission platforms. NomNom loses traffic "
    "under all paid regimes but is not itself charged by the LLM — its loss is a competitive "
    "externality from rival sponsorship."
)

subheading("5.3  Experiment 3: Sponsorship Bias Sweep")
body(
    "Experiment 3 maps the continuous relationship between bias strength (0.0–1.0) and "
    "consumer outcomes under the CPC regime."
)

caption("Table 4. Experiment 3: Bias strength sweep under CPC regime. No-LLM baseline: intent fulfilment = 0.683.")
add_table(
    ["Bias Strength", "Shortlist Relevance", "Intent Fulfilment", "LLM Revenue", "Platform Net"],
    [
        ["0.0 (neutral ranking)", "0.973", "0.747", "$5,113",  "$62,770"],
        ["0.1",                   "0.969", "0.741", "$5,787",  "$66,586"],
        ["0.2",                   "0.954", "0.731", "$6,177",  "$67,984"],
        ["0.3 (default)",         "0.929", "0.716", "$6,492",  "$69,077"],
        ["0.5",                   "0.865", "0.689", "$6,901",  "$69,444"],
        ["0.7",                   "0.816", "0.663", "$7,174",  "$70,195"],
        ["1.0 (maximum)",         "0.801", "0.649", "$7,259",  "$70,076"],
    ]
)

body(
    "Consumer harm is monotonically increasing in bias strength but LLM revenue gains are "
    "sharply diminishing above bias = 0.5. Raising bias from 0.5 to 1.0 adds only $358 (+5.2%) "
    "in LLM revenue while reducing intent fulfilment by 4.0 pp. At maximum bias (1.0), intent "
    "fulfilment (0.649) falls below the no-LLM baseline of 0.683 — the LLM with strong paid "
    "bias is worse for consumers than loyal direct platform search. A regulatory bias cap in the "
    "0.1–0.2 range would preserve nearly all consumer welfare gains at minimal LLM revenue cost."
)

subheading("5.4  Experiment 4: Consumer LLM Trust Sweep")
body(
    "Experiment 4 tests whether the LLM's effects on market structure and match quality depend "
    "on the level of consumer trust in the AI."
)

caption("Table 5. Experiment 4: Consumer LLM trust sweep.")
add_table(
    ["Trust Range", "Conv. Rate (LLM)", "Conv. Rate (No-LLM)", "HHI (LLM)", "HHI (No-LLM)", "Intent Fulfilment"],
    [
        ["Very Low (0.0–0.2)",  "78.6%", "80.4%", "0.0216", "0.0202", "0.748"],
        ["Low (0.2–0.4)",       "80.2%", "80.4%", "0.0216", "0.0202", "0.747"],
        ["Medium (0.4–0.6)",    "82.0%", "80.4%", "0.0216", "0.0202", "0.747"],
        ["High (0.6–0.8)",      "83.7%", "80.4%", "0.0215", "0.0202", "0.748"],
        ["Very High (0.8–1.0)", "85.2%", "80.4%", "0.0216", "0.0202", "0.749"],
    ]
)

body(
    "Trust has two distinct effects. A demand effect: higher trust raises conversion by "
    "approximately 6.6 pp across the full range (78.6% to 85.2%), because consumers who trust "
    "the LLM more follow its recommendations more often. A quality effect: intent fulfilment "
    "is essentially invariant across all trust bands (0.747–0.749), because the LLM's algorithm "
    "— not consumer trust — determines which offers appear in the shortlist. Concentration (HHI) "
    "is marginally higher under the LLM (0.0215–0.0216) than under loyal direct search (0.0202) "
    "at all trust levels."
)

subheading("5.5  Attribution Analysis")
body(
    "The counterfactual attribution module estimates marginal treatment effects on individual "
    "purchase probability for a 30-consumer sample with 50 Monte Carlo draws each, under the "
    "CPC regime."
)

caption("Table 6. Attribution analysis: marginal treatment effects on purchase probability (n = 30 consumers, 50 draws each).")
add_table(
    ["Effect", "Mean", "Std. Dev.", "Min", "Max"],
    [
        ["Marginal LLM effect",   "+0.014", "0.023", "−0.020", "+0.060"],
        ["Marginal bias effect",  "+0.003", "0.016", "−0.020", "+0.040"],
        ["Marginal promo effect", "+0.001", "0.018", "−0.020", "+0.060"],
    ]
)

body(
    "The LLM raises individual conversion probability by approximately 1.4 pp on average — "
    "consistent with Experiment 1's aggregate 1.2 pp gain. The bias effect is small (+0.3 pp), "
    "confirming that platforms pay primarily for order-share redistribution rather than aggregate "
    "demand creation. Promotional discounts have near-zero average marginal effect."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

heading("6. Discussion")

subheading("6.1  Interpretation")
body(
    "Experiment 1 reveals a more nuanced story than naive cross-platform aggregation would suggest. "
    "The 9.3% intent fulfilment gain under neutral LLM intermediation is real and structurally "
    "motivated — it arises from cross-platform aggregation that eliminates cuisine-match failures "
    "due to platform sampling. But the magnitude is substantially smaller than would result from "
    "an unfiltered comparison against a baseline without loyalty-based search, because the no-LLM "
    "consumers already apply a cuisine filter and check their most-loyal platforms first. The "
    "negative platform WTP (−$1,483) is the key new finding: neutral LLM intermediation creates "
    "consumer value at platform expense, by routing away from commission-maximising placements "
    "toward consumer-fit-maximising ones."
)
body(
    "Experiment 2 establishes that the payment mechanism fundamentally determines who captures "
    "value. CPFI achieves an alignment property no other regime possesses: because the LLM's "
    "marginal payment from a sponsored recommendation scales with that offer's intent fulfilment, "
    "the bias term naturally favours on-intent sponsored offers over off-intent ones. The result "
    "is a shortlist quality (0.981) that actually exceeds neutral (0.973), because CPFI "
    "incentivises the LLM to promote sponsored restaurants only when they are also good fits. "
    "This is a rare instance where incentive alignment and profit maximisation are simultaneously "
    "achievable, though at a revenue ceiling lower than CPC."
)
body(
    "Experiment 3 provides an empirical basis for regulatory thresholds. The welfare crossover — "
    "where LLM-mediated intent fulfilment falls below the no-LLM baseline — occurs between bias = "
    "0.7 and bias = 1.0. A regulatory cap in the 0.1–0.2 range would preserve approximately 95% "
    "of zero-bias welfare gains while permitting viable monetisation, mirroring disclosure-style "
    "regulation that caps paid placement weight rather than prohibiting it."
)
body(
    "Experiment 4 confirms that the welfare benefits of LLM intermediation (better intent match) "
    "do not depend on mass consumer adoption — they are determined by the algorithm, not by trust "
    "levels. However, conversion does depend strongly on trust, meaning that a nascent LLM product "
    "with low consumer trust will struggle to capture demand benefits even if its matching "
    "algorithm is sound."
)

subheading("6.2  Limitations")
body(
    "Several simplifying assumptions bound external validity. The model is a static, one-shot "
    "game: agents do not dynamically adjust strategies across rounds. In reality, restaurants "
    "would adjust promotional budgets in response to LLM rankings and platforms would compete "
    "on commission rates, potentially producing richer equilibrium behaviour including "
    "sponsorship arms races. Additional limitations include: passive restaurant agents that do "
    "not respond to ranking outcomes; a single LLM monopoly assumption that ignores competitive "
    "LLM dynamics; fixed consumer trust drawn at the start, precluding Bayesian updating as "
    "consumers observe recommendation quality over time; the hard-zero cuisine criterion in "
    "intent fulfilment, which, combined with the no-LLM baseline's cuisine filter, may narrow "
    "the measured difference between the two worlds; and no platform network effects, which "
    "would resist disintermediation in practice."
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

heading("7. Conclusion")
body(
    "This study provides simulation-based evidence that LLM intermediation in food delivery "
    "markets is moderately welfare-improving in its neutral form, generating gains in intent "
    "fulfilment (+9.3%), consumer price (−$1.71), and delivery time (−1.7 min) relative to "
    "loyalty-aware direct platform search. However, neutral LLM access generates negative "
    "platform WTP: platforms collectively lose revenue when the LLM routes for consumer fit "
    "rather than commission maximisation. Paid regimes restore positive platform WTP but "
    "differ dramatically in their consumer welfare consequences."
)
body(
    "CPFI is the recommended payment structure: its sponsorship bias is modulated by the "
    "offer's predicted intent fulfilment, aligning the LLM's financial incentive with the "
    "consumer's welfare. CPFI achieves better shortlist quality than even the neutral baseline "
    "(0.981 vs. 0.973) while generating meaningful LLM revenue and positive platform net "
    "returns. CPC produces the most severe consumer harm and extreme market concentration, "
    "with the highest-paying platform capturing nearly two-thirds of all orders."
)
body(
    "Sponsorship bias has a regulatory crossover point: at bias strength ≥ 0.7–1.0, the LLM "
    "is worse than loyal direct search on consumer intent fulfilment. A bias cap in the "
    "0.1–0.2 range would preserve nearly all consumer welfare gains. The market concentration "
    "effects of LLM intermediation are structurally small under neutral routing but become "
    "severe under CPC. Trust adoption does not affect match quality, only demand volume."
)
body(
    "Taken together, the findings suggest that LLM-mediated food delivery markets are "
    "welfare-improving if governed by CPFI-style payment structures and bias caps, and "
    "welfare-degrading if governed by CPC structures transplanted from traditional digital "
    "advertising without modification."
)

heading("Future Research")
body(
    "Priority extensions include: (1) a dynamic repeated game with strategy-updating restaurant "
    "and platform agents to study Nash equilibrium convergence and sponsorship arms races; "
    "(2) multi-LLM competition modelling where multiple AI assistants compete for platform "
    "sponsorship revenue; (3) real embedding-based semantic cuisine matching via "
    "sentence-transformer cosine similarity, replacing the hand-coded 9×9 table; "
    "(4) Bayesian consumer trust dynamics where users update trust based on observed "
    "recommendation quality; (5) restaurant-side LLM sponsorship as an independent payment "
    "layer; and (6) a CPWI (cost-per-welfare-improvement) regime aligning LLM incentives "
    "with utility generation."
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

heading("References")

refs = [
    "Anderson, S. P., de Palma, A., & Thisse, J. F. (1992). Discrete choice theory of product differentiation. MIT Press.",
    "Bergemann, D., & Välimäki, J. (2019). Dynamic mechanism design: An introduction. Journal of Economic Literature, 57(2), 235–274.",
    "Eisenmann, T., Parker, G., & Van Alstyne, M. (2006). Strategies for two-sided markets. Harvard Business Review, 84(10), 92–101.",
    "Hagiu, A., & Wright, J. (2015). Multi-sided platforms. International Journal of Industrial Organization, 43, 162–174.",
    "McFadden, D. (1974). Conditional logit analysis of qualitative choice behaviour. In P. Zarembka (Ed.), Frontiers in Econometrics (pp. 105–142). Academic Press.",
    "Rochet, J. C., & Tirole, J. (2003). Platform competition in two-sided markets. Journal of the European Economic Association, 1(4), 990–1029.",
    "Varian, H. R. (2007). Position auctions. International Journal of Industrial Organization, 25(6), 1163–1178.",
    "Wei, J., et al. (2022). Emergent abilities of large language models. Transactions on Machine Learning Research.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    set_font(run, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("report_final.docx")
print("Saved report_final.docx")
